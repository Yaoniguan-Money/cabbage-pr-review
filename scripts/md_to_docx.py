"""将 Markdown 审阅报告转换为 Word (.docx)。"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.oxml import OxmlElement


def _set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def _add_formatted_run(paragraph, text: str, *, bold: bool = False, code: bool = False) -> None:
    """解析 **bold** 与 `code` 内联格式。"""
    pattern = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        chunk = match.group(0)
        if chunk.startswith("**"):
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
        else:
            run = paragraph.add_run(chunk[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])
    if not text and bold:
        paragraph.add_run("").bold = True


def _parse_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [cell.strip() for cell in line.split("|")]


def _is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\|[\s\-:|]+\|\s*$", line.strip()))


def convert_md_to_docx(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    # 页面边距
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.font.size = Pt(10.5)

    i = 0
    in_code = False
    code_lang = ""
    code_lines: list[str] = []
    in_mermaid = False

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 代码块
        if stripped.startswith("```"):
            if not in_code:
                in_code = True
                code_lang = stripped[3:].strip()
                in_mermaid = code_lang == "mermaid"
                code_lines = []
            else:
                in_code = False
                if in_mermaid:
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(6)
                    run = p.add_run("【架构图 / 流程图（Mermaid 源码见 Markdown 版）】")
                    run.bold = True
                    run.font.color.rgb = RGBColor(0x1A, 0x56, 0x8E)
                    for cl in code_lines[:12]:
                        cp = doc.add_paragraph(cl)
                        cp.paragraph_format.left_indent = Cm(0.5)
                        for r in cp.runs:
                            r.font.name = "Consolas"
                            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
                            r.font.size = Pt(8)
                    if len(code_lines) > 12:
                        doc.add_paragraph(f"… 共 {len(code_lines)} 行，完整内容见 docs/RULES_ENGINE_REVIEW.md")
                    in_mermaid = False
                else:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(0.4)
                    p.paragraph_format.space_before = Pt(4)
                    p.paragraph_format.space_after = Pt(4)
                    pf = p.paragraph_format
                    pPr = p._element.get_or_add_pPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:fill"), "F5F5F5")
                    pPr.append(shd)
                    run = p.add_run("\n".join(code_lines))
                    run.font.name = "Consolas"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
                    run.font.size = Pt(9)
                code_lines = []
                code_lang = ""
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # 跳过分隔线
        if stripped in {"---", "***", "___"}:
            i += 1
            continue

        # 空行
        if not stripped:
            i += 1
            continue

        # 表格
        if stripped.startswith("|") and i + 1 < len(lines) and _is_table_separator(lines[i + 1]):
            headers = _parse_table_row(line)
            i += 2
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(_parse_table_row(lines[i]))
                i += 1
            col_count = len(headers)
            table = doc.add_table(rows=1 + len(rows), cols=col_count)
            table.style = "Table Grid"
            for ci, h in enumerate(headers):
                cell = table.rows[0].cells[ci]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(h.replace("**", ""))
                run.bold = True
                _set_cell_shading(cell, "E8EEF4")
            for ri, row in enumerate(rows):
                for ci in range(col_count):
                    val = row[ci] if ci < len(row) else ""
                    cell = table.rows[ri + 1].cells[ci]
                    cell.text = ""
                    _add_formatted_run(cell.paragraphs[0], val.replace("**", ""))
            doc.add_paragraph()
            continue

        # 标题
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
            if level == 1:
                p = doc.add_heading(text, level=0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc.add_heading(text, level=min(level, 3))
            i += 1
            continue

        # 引用块
        if stripped.startswith(">"):
            text = stripped.lstrip("> ").strip()
            text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.8)
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # 无序列表
        if stripped.startswith("- ") or stripped.startswith("* "):
            text = stripped[2:].strip()
            text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_run(p, text)
            i += 1
            continue

        # 有序列表
        ol_match = re.match(r"^(\d+)\.\s+(.*)", stripped)
        if ol_match:
            text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", ol_match.group(2))
            p = doc.add_paragraph(style="List Number")
            _add_formatted_run(p, text)
            i += 1
            continue

        # 普通段落
        text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", stripped)
        p = doc.add_paragraph()
        _add_formatted_run(p, text)
        i += 1

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(docx_path))
    print(f"已生成: {docx_path}")


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    md = root / "docs" / "RULES_ENGINE_REVIEW.md"
    out = root / "docs" / "RULES_ENGINE_REVIEW.docx"
    if len(sys.argv) >= 2:
        md = Path(sys.argv[1])
    if len(sys.argv) >= 3:
        out = Path(sys.argv[2])
    convert_md_to_docx(md, out)
